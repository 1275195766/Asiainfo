import docx
import pandas as pd
import random

wordpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\规格1.docx'
excelpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\清单.xlsx'
# lb_path=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\类别.xlsx'
df=pd.read_excel(excelpath)
doc=docx.Document(wordpath)
# print(df.shape)
A1='天天基金,蚂蚁财富,掌上基金,现金宝,滚雪球基金,同花顺爱基金,e钱包,嘉实理财嘉,南方基金,华夏基金管家,广发基金,天虹爱理财,利得基金,中欧钱滚滚,银华生利宝,富国富钱包,博时基金,招商招钱宝,工银现金快线,鹏华A加钱包,建信基金,中银基金,安信基金,兴全基金,国泰基金,国投瑞银,大成基金,景顺长城基金,交银基金,基金E站,广大保德信基金,融通基金,华安基金,华商基金,诺安基金,东方基金,华泰柏瑞,长信基金,农银汇理基金,长盛基金,贝壳APP'.split(',')
A2='ok车险,保险大咖,保险驾到,快保,平安保险商城,i云保,平安健康,水滴保险,梧桐树保险,向日葵宝信啊,小贝保险,小智保险,阳光保险,阳光惠生活,最惠保,太平洋保险,平安金管家,平安车主,国寿e宝,神行太保,众安,中国人保,泰康在线,平安健康,大象保险,唯彩看球,付临门,还呗,好分期,来分期,趣花分期,时光分期,我爱卡,平安证券,华泰证券,财富证券,智远一户通,国泰君安,银河证券,广发证券,海通证券'.split(',')

B1='360贷款,51人品贷,安逸花,飞贷,给你花,好易借,恒易贷,及贷,极速贷款,急现贷,捷信金融,借贷宝,借乐花,快贷,陆金所,魔法现金,你我贷,拍拍贷,平安贷款,平安好贷,人人货,闪贷,省呗,微贷,维信卡卡贷,小米贷款,小赢卡贷,现金巴士,新橙优品,信而富,有钱花,招联金融,国美易卡,海尔消费金融,玖富万卡,钱伴,钱站,支付宝,有信钱包,蛋卷基金,银收宝'.split(',')
B2='国金证券,方正证券,长江证券,招商证券,涨乐财富通,小方,金太阳,申万宏源大赢家,长江e号,安信手机证券,中心建投证券,掌中投,金阳光移动证券,华彩人生新版,中银国际,悟空理财,陆金所,理财通,招联金融,橘子理财,朵朵金融,金融凤凰,财迷之家,国美金融,好兴动,京东金融,京东钱包,马上金融,平安普惠,爱钱进,融360,苏宁金融,苏宁消费金融,微众银行,子主题19,沃钱包,小花钱包,翼支付,云闪付,中信证券,淘宝'.split(',')


def set_head():
    i = 0
    while True:
        if doc.paragraphs[i].style.name == 'Heading 4':

            # print('i', i)
            para_text = doc.paragraphs[i].text
            print(para_text)


            doc.paragraphs[i+1].insert_paragraph_before(text='功能描述',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='外部信息',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='内部信息',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='数据处理',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='系统输出',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='查询',style='Heading 5')
            doc.paragraphs[i+1].insert_paragraph_before(text='生产内部数据',style='Heading 5')


        length = len(doc.paragraphs)
        if i+1 == length:
            break
        else:
            i += 1


def get_index(start):
    for i in range(start+1,len(doc.paragraphs)):

            if doc.paragraphs[i].style.name == 'Heading 4':
                print('get_index',doc.paragraphs[i].text)

                if len(df.loc[df.客户需求 == doc.paragraphs[i].text].index)==0:
                    print('index=0',df.loc[df.客户需求 == doc.paragraphs[i].text].index)
                    end=0
                else:
                # try:
                    end=df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]
                # except BaseException as e1:
                #     print('e1',e1)

                return end


def rotate(nums, k):
    start = 0
    end = len(nums) - 1
    k = k % len(nums)

    reverse(nums, start, end)
    reverse(nums, start, k - 1)
    reverse(nums, k, end)

    return nums


def reverse(nums, start, end):
    while start < end:
        t = nums[start]
        nums[start] = nums[end]
        nums[end] = t
        # print(nums)
        start += 1
        end -= 1


nbxx_list=['用户资料信息日表','基站信息维度表','经纬度信息维度表','用户终端品牌信息表']
xtsc_list=['伪基站活动区域GIS监控结果表','自定义监控区域管理信息表','监控区域告警信息展示信息表']
cx_list=['对伪基站活动轨迹查询','对今日区域轨迹列表进行查询','对昨日影响小区数Top10伪基站轨迹进行查询']

def write_para(string,j):
    doc.paragraphs[j + 1].insert_paragraph_before(text=string,style=None)


i=0
n=0
while True:
    # print('开始写入')
    try:
        if doc.paragraphs[i].style.name=='Heading 3':
            head3_text = doc.paragraphs[i].text

        if doc.paragraphs[i].style.name=='Heading 4':
            para_text = doc.paragraphs[i].text
            # print('Heading 4  i', i)
            print('head3_text', head3_text)
            print('Heading 4',para_text,'i',i)

            # print(df.loc[df.客户需求 == doc.paragraphs[i].text,'功能过程'].values[0])
            # print(df.loc[df.客户需求 == doc.paragraphs[i].text].index[0])
            start = df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]
            end = get_index(i)
            if end==0:
                i+=1
                continue
            xh=df.loc[df.客户需求 == para_text, '编号'].values[0]
            print('编号',xh,'start',start,'end',end)
            j=i
            while True :

                try:
                    if doc.paragraphs[j].text=='功能描述':
                        print('功能描述')
                        doc.paragraphs[j + 1].insert_paragraph_before(text=df.loc[df.客户需求 == para_text, '功能过程'].values[0], style=None)
                    elif doc.paragraphs[j].text=='外部信息':
                        doc.paragraphs[j + 1].insert_paragraph_before(text='用户访问{0}{1}\n{2}{3}'.format(A1[n],A2[n],B1[n],B2[n]),style=None)


                    elif doc.paragraphs[j].text=='内部信息':
                        # write_para('状态信息表',j)
                        # doc.paragraphs[j + 1].insert_paragraph_before(text='状态信息表\n用户资料信息日表\n基站信息维度表\n经纬度信息维度表\n用户终端品牌信息表', style=None)

                        doc.paragraphs[j + 1].insert_paragraph_before(text='{0}的{1}的数据表\n{1}数据'.format(head3_text,para_text), style=None)


                    elif doc.paragraphs[j].text=='数据处理':

                        if end==0:
                            break
                        print('数据处理',df.iloc[start:end, 6].values)
                        doc.paragraphs[j+1].insert_paragraph_before(text='\n'.join(df.iloc[start:end, 6].values),style=None)
                    elif doc.paragraphs[j].text=='系统输出':
                        doc.paragraphs[j + 1].insert_paragraph_before(text='{0}表\n{1}{0}数据'.format(para_text,head3_text), style=None)

                    elif doc.paragraphs[j].text=='查询':
                        doc.paragraphs[j + 1].insert_paragraph_before(text=random.choice(cx_list), style=None)


                    elif doc.paragraphs[j].text=='生产内部数据':
                        # doc.paragraphs[j + 1].insert_paragraph_before(text=df.loc[df.客户需求 == para_text, '数据属性'].values[0],style=None)
                        A1_A2=A1[n]+A2[n]
                        B1_B2=B1[n]+B2[n]
                        doc.paragraphs[j + 1].insert_paragraph_before(text='用户访问{0}的数据\n{1}的{2}表的数据\n用户访问{3}的数据'.format(A1_A2,head3_text,para_text,B1_B2),style=None)
                        df = df.drop(labels=range(start, end), axis=0).reset_index(drop=True)
                        # print('删除')
                        if n==40:
                            rotate(A2,1)
                            rotate(B2,1)
                            n=0
                        break


                except BaseException as e1:
                    print('内层循环',e1)
                    df = df.drop(labels=range(start, end), axis=0).reset_index(drop=True)
                    break
                finally:

                    j+=1
                pass
        # doc.save(r'F:\pythonPrj\Asiainfo\word\规格1.docx')


                # doc.save(r'F:\pythonPrj\Asiainfo\word\返回图标1.docx')
        length=len(doc.paragraphs)
        # print('length:'+str(length)+'   i  '+str(i))
        if i+1 == length:
            break
        else:

            i += 1
            # print('i:',i)
        pass
    except BaseException as e:
        print('外层循环',e)
        with open(r'F:\pythonPrj\Asiainfo\word\规格1.log','a+') as w:
            w.write(para_text+'\n')
        if i+1 == length:
            break
        else:
            i += 1
        pass




pass
doc.save(r'F:\pythonPrj\Asiainfo\word\规格1.docx')
# print(doc.paragraphs[6].text)