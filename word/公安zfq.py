import docx
import pandas as pd

wordpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\zfq.docx'
excelpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\zfq.xlsx'
df=pd.read_excel(excelpath)
doc=docx.Document(wordpath)
# for i in range(len(doc.paragraphs)):
#     if doc.paragraphs[i].style.name == 'Heading 4':
#         print(doc.paragraphs[i].text)
def get_index(start):
    for i in range(start+1,len(doc.paragraphs)):

            if doc.paragraphs[i].style.name == 'Heading 4':
                print('get_index',doc.paragraphs[i].text)
                if len(df.loc[df.客户需求 == doc.paragraphs[i].text].index)==0:
                    print('index=0',df.loc[df.客户需求 == doc.paragraphs[i].text].index)
                    end=0
                else:
                # try:
                    end=df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]
                # except BaseException as e1:
                #     print('e1',e1)

                return end

def write_para(string,j):
    doc.paragraphs[j + 1].insert_paragraph_before(text=string,style=None)

# for i in range(len(doc.paragraphs)):
i=0

while True:
    # print('iii',i,doc.paragraphs[i].text)
    try:
        if doc.paragraphs[i].style.name=='Heading 4':
            print('i',i)
            para_text=doc.paragraphs[i].text
            print(para_text)

            # print(df.loc[df.客户需求 == doc.paragraphs[i].text,'功能过程'].values[0])
            # print(df.loc[df.客户需求 == doc.paragraphs[i].text].index[0])
            start = df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]
            end = get_index(i)
            if end==0:
                i+=1
                continue
            print('start',start,'end',end)
            j=i
            while True :
                # print('j',j,doc.paragraphs[j].text)
                # print('状态' in para_text)
                # print('日期'  in para_text )
                # print(('主题' or '时间' or '日期') in para_text)
                try:
                    if doc.paragraphs[j].text=='功能描述':
                        print('功能描述')
                        doc.paragraphs[j + 1].insert_paragraph_before(text=df.loc[df.客户需求 == para_text, '功能过程'].values[0], style=None)
                    elif doc.paragraphs[j].text=='外部信息':
                        doc.paragraphs[j + 1].insert_paragraph_before(text='其他接口推送的数据',
                                                                      style=None)


                    elif doc.paragraphs[j].text=='内部信息':
                        # write_para('状态信息表',j)
                        doc.paragraphs[j + 1].insert_paragraph_before(text='状态信息表\n用户资料信息日表\n基站信息维度表\n经纬度信息维度表\n用户终端品牌信息表', style=None)


                    elif doc.paragraphs[j].text=='数据处理':

                        if end==0:
                            break
                        print('数据处理',df.iloc[start:end, 7].values)
                        doc.paragraphs[j+1].insert_paragraph_before(text='\n'.join(df.iloc[start:end, 7].values),style=None)
                    elif doc.paragraphs[j].text=='系统输出':
                        print('系统输出')
                        doc.paragraphs[j + 1].insert_paragraph_before(text='状态信息输出表\n伪基站活动区域GIS监控结果表\n自定义监控区域管理信息表\n监控区域告警信息展示信息表', style=None)

                    elif doc.paragraphs[j].text=='查询':

                        doc.paragraphs[j + 1].insert_paragraph_before(text='对伪基站活动轨迹查询\n对今日区域轨迹列表进行查询\n对昨日影响小区数Top10伪基站轨迹进行查询', style=None)


                    elif doc.paragraphs[j].text=='生产内部数据':
                        doc.paragraphs[j + 1].insert_paragraph_before(text=df.loc[df.客户需求 == para_text, '数据属性'].values[0],
                                                                  style=None)
                        df = df.drop(labels=range(start, end), axis=0).reset_index(drop=True)
                        print('删除')
                        break


                except BaseException as e1:
                    print('e1',e1)
                    df = df.drop(labels=range(start, end), axis=0).reset_index(drop=True)
                    break
                finally:

                    j+=1
                pass
            # pass
            doc.save(r'F:\pythonPrj\Asiainfo\word\zfq.docx')

        # doc.save(r'F:\pythonPrj\Asiainfo\word\返回图标1.docx')
        length=len(doc.paragraphs)
        # print('length:'+str(length)+'   i  '+str(i))
        if i+1 == length:
            break
        else:
            i += 1
        pass
    except BaseException as e:
        print('e',e)
        if i+1 == length:
            break
        else:
            i += 1
        pass





# print(doc.paragraphs[6].text)