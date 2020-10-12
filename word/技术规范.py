import docx
import pandas as pd

wordpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\关系分析.docx'
excelpath=r'D:\文档\OneDrive-wyx\OneDrive - wuyaoxue\亚信实习\工作\公安三期项目\工作簿1.xlsx'

df=pd.read_excel(excelpath)
doc=docx.Document(wordpath)


def get_index(start):
    for i in range(start+1,len(doc.paragraphs)):

            if doc.paragraphs[i].style.name == 'Heading 4':
                print('get_index',doc.paragraphs[i].text)
                if len(df.loc[df.客户需求 == doc.paragraphs[i].text].index)==0:
                    print('index=0',df.loc[df.客户需求 == doc.paragraphs[i].text].index)
                    end=0
                else:

                    end=df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]

                return end


i=0

while True:
    try:
        if doc.paragraphs[i].style.name=='Heading 4':
            print('i',i)
            para_text=doc.paragraphs[i].text
            print(para_text)

            start = df.loc[df.客户需求 == doc.paragraphs[i].text].index[0]
            end = get_index(i)
            if end==0:
                i+=1
                continue
            print('start',start,'end',end)
            if end == 0:
                break
            print('数据处理', df.iloc[start:end, 7].values)

            j=i
            doc.paragraphs[j + 1].insert_paragraph_before(text='\n'.join(df.iloc[start:end, 7].values), style=None)
            df = df.drop(labels=range(start, end), axis=0).reset_index(drop=True)
            doc.save(r'F:\pythonPrj\Asiainfo\word\关系分析.docx')
        length=len(doc.paragraphs)
        # print('length:'+str(length)+'   i  '+str(i))
        if i+1 == length:
            break
        else:
            i += 1
        pass
    except BaseException as e:
        print('e',e)
        if i+1 == length:
            break
        else:
            i += 1
        pass




pass
doc.save(r'F:\pythonPrj\Asiainfo\word\关系分析.docx')